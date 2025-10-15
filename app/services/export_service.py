from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session
from app.models.rfp_project import RFPProject
from app.models.rfp_question import RFPQuestion
from typing import List
import io
from datetime import datetime

class ExportService:
    
    @staticmethod
    def get_trust_color(trust_score: float) -> tuple:
        if trust_score >= 80:
            return (144, 238, 144)
        elif trust_score >= 50:
            return (255, 255, 153)
        else:
            return (255, 182, 193)
    
    @staticmethod
    def export_to_excel(rfp: RFPProject, questions: List[RFPQuestion]) -> bytes:
        wb = Workbook()
        ws = wb.active
        ws.title = "RFP Responses"
        
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)
        
        headers = ["Question", "Answer", "Trust Score", "Source Type"]
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        for idx, q in enumerate(questions, start=2):
            ws.cell(row=idx, column=1, value=q.question_text)
            ws.cell(row=idx, column=2, value=q.answer_text)
            
            trust_cell = ws.cell(row=idx, column=3, value=round(q.trust_score, 2))
            trust_cell.alignment = Alignment(horizontal="center")
            
            color = ExportService.get_trust_color(q.trust_score)
            trust_cell.fill = PatternFill(
                start_color=f"{color[0]:02x}{color[1]:02x}{color[2]:02x}",
                end_color=f"{color[0]:02x}{color[1]:02x}{color[2]:02x}",
                fill_type="solid"
            )
            
            ws.cell(row=idx, column=4, value=q.source_type)
            
            for col in range(1, 5):
                ws.cell(row=idx, column=col).alignment = Alignment(wrap_text=True, vertical="top")
        
        metadata_ws = wb.create_sheet("Metadata")
        metadata = [
            ["RFP Name", rfp.rfp_name],
            ["Generated Date", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")],
            ["Total Questions", len(questions)],
            ["Average Trust Score", round(sum(q.trust_score for q in questions) / len(questions), 2) if questions else 0]
        ]
        
        for row_idx, (key, value) in enumerate(metadata, start=1):
            metadata_ws.cell(row=row_idx, column=1, value=key).font = Font(bold=True)
            metadata_ws.cell(row=row_idx, column=2, value=value)
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def export_to_docx(rfp: RFPProject, questions: List[RFPQuestion]) -> bytes:
        doc = Document()
        
        title = doc.add_heading(rfp.rfp_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph("")
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "Question"
        header_cells[1].text = "Answer"
        header_cells[2].text = "Trust Score"
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for q in questions:
            row = table.add_row()
            row.cells[0].text = q.question_text
            row.cells[1].text = q.answer_text
            
            trust_text = f"{round(q.trust_score, 2)}%"
            if q.trust_score >= 80:
                trust_text += " ✓"
            elif q.trust_score < 50:
                trust_text += " ⚠"
            
            row.cells[2].text = trust_text
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def export_to_pdf(rfp: RFPProject, questions: List[RFPQuestion]) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#366092'),
            spaceAfter=30,
            alignment=1
        )
        
        elements.append(Paragraph(rfp.rfp_name, title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        metadata_text = f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}<br/>Total Questions: {len(questions)}"
        elements.append(Paragraph(metadata_text, styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
        
        data = [['Question', 'Answer', 'Trust']]
        
        for q in questions:
            data.append([
                Paragraph(q.question_text, styles['Normal']),
                Paragraph(q.answer_text[:500] + "..." if len(q.answer_text) > 500 else q.answer_text, styles['Normal']),
                f"{round(q.trust_score, 1)}%"
            ])
        
        table = Table(data, colWidths=[2*inch, 4*inch, 0.8*inch])
        
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ])
        
        for idx, q in enumerate(questions, start=1):
            color = ExportService.get_trust_color(q.trust_score)
            table_style.add('BACKGROUND', (2, idx), (2, idx), colors.Color(color[0]/255, color[1]/255, color[2]/255))
        
        table.setStyle(table_style)
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()