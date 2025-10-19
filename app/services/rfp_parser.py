from typing import List, Dict
import json
import logging
from app.services.llm_factory import LLMFactory
from app.prompts.question_extractor import (
    EXTRACT_QUESTIONS_SYSTEM,
    EXTRACT_QUESTIONS_USER,
)
from langchain_core.prompts import ChatPromptTemplate
import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class RFPParser:
    
    QUESTION_PATTERNS = [
        r'^\d+[\.\)]\s+(.+)',
        r'^[A-Z]\.\s+(.+)',
        r'^[a-z][\.\)]\s+(.+)',
        r'^\(?[ivx]+\)?\s+(.+)',
        r'^Q\d+[:\.\)]\s*(.+)',
        r'^Question\s+\d+[:\.\)]\s*(.+)',
    ]
    
    QUESTION_KEYWORDS = [
        'describe', 'explain', 'what', 'how', 'why', 'when', 'where', 'who',
        'do you', 'does your', 'can you', 'will you', 'have you', 'are you',
        'is your', 'provide', 'list', 'detail', 'outline', 'specify',
        'demonstrate', 'confirm', 'identify', 'indicate', 'please',
        'tell us', 'give us', 'show us', 'supply', 'submit'
    ]
    
    @staticmethod
    def extract_questions(file_path: str, file_type: str) -> List[str]:
        try:
            if file_type.endswith('.pdf'):
                text = RFPParser._extract_pdf_text(file_path)
            elif file_type.endswith('.docx'):
                text = RFPParser._extract_docx_text(file_path)
            elif file_type.endswith(('.txt', '.md')):
                text = RFPParser._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            questions = RFPParser._ai_extract_questions(text)
            
            if not questions or len(questions) == 0:
                logger.warning("AI extraction returned no questions, falling back to regex")
                questions = RFPParser._fallback_extract(text)
            
            return questions
            
        except Exception as e:
            logger.error(f"Error in AI extraction: {str(e)}, using fallback")
            return RFPParser._fallback_extract(text)
    
    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")
    
    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
    
    @staticmethod
    def _extract_txt_text(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text: {str(e)}")
    
    @staticmethod
    def _ai_extract_questions(text: str) -> List[str]:
        try:
            llm = LLMFactory.get_llm("gemini-2.0-flash", temperature=0)
            
            prompt = ChatPromptTemplate.from_messages([
                ("system", EXTRACT_QUESTIONS_SYSTEM),
                ("human", EXTRACT_QUESTIONS_USER)
            ])
            
            chain = prompt | llm
            
            if len(text) > 50000:
                text = text[:50000]
            
            response = chain.invoke({"text": text})
            
            content = response.content.strip()
            if content.startswith('```json'):
                content = content[7:]
            if content.startswith('```'):
                content = content[3:]
            if content.endswith('```'):
                content = content[:-3]
            content = content.strip()
            
            questions_data = json.loads(content)
            
            questions = [q["text"] for q in questions_data if q.get("text")]
            
            logger.info(f"AI extracted {len(questions)} questions")
            return questions
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error in AI extraction: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"AI extraction error: {str(e)}")
            raise
    
    @staticmethod
    def _fallback_extract(text: str) -> List[str]:
        import re
        
        questions = []
        lines = text.split('\n')
        
        current_question = None
        question_number = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line or len(line) < 10:
                continue
            
            is_question = False
            question_text = line
            
            if line.endswith('?'):
                is_question = True
            
            for pattern in RFPParser.QUESTION_PATTERNS:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    is_question = True
                    if match.groups():
                        question_text = match.group(1).strip()
                    break
            
            if not is_question:
                line_lower = line.lower()
                if any(line_lower.startswith(keyword) for keyword in RFPParser.QUESTION_KEYWORDS):
                    is_question = True
            
            if is_question:
                if current_question and len(current_question) < 500:
                    questions.append(current_question)
                
                question_number += 1
                current_question = question_text
            elif current_question and len(line) > 20:
                if len(current_question) < 300:
                    current_question += ' ' + line
        
        if current_question:
            questions.append(current_question)
        
        filtered_questions = [q for q in questions if len(q.strip()) >= 15]
        
        logger.info(f"Fallback extracted {len(filtered_questions)} questions")
        return filtered_questions